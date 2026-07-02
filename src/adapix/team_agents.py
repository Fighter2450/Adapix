"""AI Team — specialist agent personalities for business owners.

Loads agent .md files from the agency-agents repo and provides per-agent
chat sessions with tool use: web search, customer data, metrics, message
drafting, and document creation.
"""
from __future__ import annotations

import json
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from anthropic import Anthropic

from .config import Settings
from .practice import load_profile


# ---------------------------------------------------------------------------
# Repo location
# ---------------------------------------------------------------------------
def _agents_repo() -> Path:
    """Where the agent definition .md files live.

    Priority: explicit env override → the copy BUNDLED inside the app
    (src/adapix/agents_library — this is what production uses; a Desktop
    folder obviously doesn't exist in a deployed container) → the legacy
    dev location on the Desktop."""
    env = os.environ.get("AGENCY_AGENTS_DIR")
    if env:
        return Path(env)
    bundled = Path(__file__).parent / "agents_library"
    if bundled.exists():
        return bundled
    return Path.home() / "Desktop" / "agency-agents"


# ---------------------------------------------------------------------------
# Curated list: slug → display category
# ---------------------------------------------------------------------------
CURATED: dict[str, str] = {
    # Strategy & Leadership
    "business-strategist":                "Strategy & Leadership",
    "specialized-chief-of-staff":         "Strategy & Leadership",
    "operations-manager":                 "Strategy & Leadership",
    "specialized-pricing-analyst":        "Strategy & Leadership",
    "change-management-consultant":       "Strategy & Leadership",
    # Marketing
    "marketing-content-creator":          "Marketing",
    "marketing-email-strategist":         "Marketing",
    "marketing-growth-hacker":            "Marketing",
    "marketing-seo-specialist":           "Marketing",
    "marketing-social-media-strategist":  "Marketing",
    "marketing-linkedin-content-creator": "Marketing",
    "marketing-instagram-curator":        "Marketing",
    "marketing-tiktok-strategist":        "Marketing",
    "marketing-pr-communications-manager":"Marketing",
    "marketing-video-optimization-specialist": "Marketing",
    "marketing-global-podcast-strategist":"Marketing",
    "marketing-app-store-optimizer":      "Marketing",
    "marketing-reddit-community-builder": "Marketing",
    # Sales
    "sales-outbound-strategist":          "Sales",
    "sales-coach":                        "Sales",
    "sales-deal-strategist":              "Sales",
    "sales-discovery-coach":              "Sales",
    "sales-outreach":                     "Sales",
    "sales-pipeline-analyst":             "Sales",
    "sales-proposal-strategist":          "Sales",
    "sales-offer-lead-gen-strategist":    "Sales",
    "sales-account-strategist":           "Sales",
    # Finance
    "chief-financial-officer":            "Finance",
    "finance-financial-analyst":          "Finance",
    "finance-bookkeeper-controller":      "Finance",
    "finance-fpa-analyst":                "Finance",
    "finance-tax-strategist":             "Finance",
    # Customer
    "customer-success-manager":           "Customer",
    "customer-service":                   "Customer",
    "healthcare-customer-service":        "Customer",
    # HR & People
    "hr-onboarding":                      "HR & People",
    "recruitment-specialist":             "HR & People",
    "organizational-psychologist":        "HR & People",
    # Legal
    "legal-compliance-checker":           "Legal",
    "legal-document-review":              "Legal",
    "data-privacy-officer":               "Legal",
    # Brand & Design
    "design-brand-guardian":              "Brand & Design",
    "design-ux-researcher":               "Brand & Design",
    "design-ui-designer":                 "Brand & Design",
    # Product
    "product-manager":                    "Product",
    "product-feedback-synthesizer":       "Product",
    "product-sprint-prioritizer":         "Product",
    "product-trend-researcher":           "Product",
    # Operations
    "supply-chain-strategist":            "Operations",
    "inventory-fulfillment-manager":      "Operations",
    # Engineering
    "engineering-senior-developer":       "Engineering",
    "engineering-backend-architect":      "Engineering",
    "engineering-frontend-developer":     "Engineering",
    "engineering-software-architect":     "Engineering",
    "engineering-prompt-engineer":        "Engineering",
    "engineering-code-reviewer":          "Engineering",
    "engineering-ai-engineer":            "Engineering",
    "testing-reality-checker":            "Engineering",
}

CATEGORY_ORDER = [
    "Strategy & Leadership",
    "Marketing",
    "Sales",
    "Finance",
    "Customer",
    "HR & People",
    "Legal",
    "Brand & Design",
    "Product",
    "Operations",
    "Engineering",
]


# ---------------------------------------------------------------------------
# Tool definitions (passed to Claude as function-calling tools)
# ---------------------------------------------------------------------------
AGENT_TOOLS = [
    {
        "name": "search_web",
        "description": (
            "Search the web for current information — competitor research, market trends, "
            "industry news, pricing benchmarks, regulatory updates, or anything requiring "
            "up-to-date knowledge beyond your training data."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "The search query. Be specific for better results.",
                }
            },
            "required": ["query"],
        },
    },
    {
        "name": "get_customers",
        "description": (
            "Query the practice's customer/patient database. Returns names, status, "
            "preferred communication channel, and treatment info. Use this when you "
            "need to know who's in the system, their current status, or to identify "
            "specific customer segments to act on."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "status": {
                    "type": "string",
                    "description": (
                        "Filter by status. Options: 'consulted_not_started', "
                        "'treatment_started', 'explicitly_declined', 'paused'. "
                        "Leave empty to get all customers."
                    ),
                },
                "limit": {
                    "type": "integer",
                    "description": "Max customers to return (default 30, max 100).",
                },
            },
        },
    },
    {
        "name": "get_metrics",
        "description": (
            "Get live business metrics from the dashboard: messages sent today, "
            "pending approval count, open escalations, and recent activity."
        ),
        "input_schema": {
            "type": "object",
            "properties": {},
        },
    },
    {
        "name": "draft_to_queue",
        "description": (
            "Draft a message for a specific customer and add it to the approval queue. "
            "The practice owner reviews and approves before it sends — nothing goes out "
            "automatically. Use this when the owner asks you to write a message for a customer."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "customer_name": {
                    "type": "string",
                    "description": "First and last name of the customer (must match a name in the database).",
                },
                "channel": {
                    "type": "string",
                    "enum": ["sms", "email"],
                    "description": "Delivery channel.",
                },
                "subject": {
                    "type": "string",
                    "description": "Email subject line (optional for SMS).",
                },
                "body": {
                    "type": "string",
                    "description": "The full message body to send to the customer.",
                },
            },
            "required": ["customer_name", "channel", "body"],
        },
    },
    {
        "name": "create_document",
        "description": (
            "Create and save a Word document (.docx) with the provided content. "
            "Use this when the owner asks for a deliverable they can download and keep — "
            "a strategy document, report, template, proposal, checklist, etc."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Document title (also used as the filename).",
                },
                "content": {
                    "type": "string",
                    "description": "Full document content in plain text or Markdown.",
                },
            },
            "required": ["title", "content"],
        },
    },
]

TOOL_LABELS = {
    "search_web":     "🔍 Web search",
    "get_customers":  "👥 Customer data",
    "get_metrics":    "📊 Business metrics",
    "draft_to_queue": "✉ Drafted to queue",
    "create_document":"📄 Document created",
}


# ---------------------------------------------------------------------------
# Tool execution
# ---------------------------------------------------------------------------
def _execute_tool(name: str, inputs: dict, settings: Settings) -> str:
    if name == "search_web":
        return _tool_search_web(inputs.get("query", ""))

    if name == "get_customers":
        return _tool_get_customers(inputs, settings)

    if name == "get_metrics":
        return _tool_get_metrics(settings)

    if name == "draft_to_queue":
        return _tool_draft_to_queue(inputs, settings)

    if name == "create_document":
        return _tool_create_document(inputs)

    return f"Unknown tool: {name}"


def _tool_search_web(query: str) -> str:
    try:
        from ddgs import DDGS
        with DDGS() as d:
            hits = list(d.text(query, max_results=5))
        if not hits:
            return f"No results found for '{query}'."
        parts = []
        for h in hits:
            title = h.get("title", "")
            body = h.get("body", "")
            url = h.get("href", "")
            parts.append(f"**{title}**\n{body}\n{url}")
        return "\n\n".join(parts)
    except Exception as e:
        return f"Search failed: {e}. Use training knowledge to answer."


def _tool_get_customers(inputs: dict, settings: Settings) -> str:
    from .db import get_session
    from .models import Patient
    try:
        limit = min(int(inputs.get("limit", 30)), 100)
        status_filter = inputs.get("status")
        with get_session(settings) as session:
            q = session.query(Patient)
            if status_filter:
                q = q.filter(Patient.status == status_filter)
            patients = q.limit(limit).all()
            if not patients:
                return "No customers found" + (f" with status '{status_filter}'" if status_filter else "") + "."
            rows = []
            for p in patients:
                name = f"{p.first_name} {p.last_name}"
                channel = p.preferred_channel or "sms"
                status = p.status or "unknown"
                treatment = p.treatment_type or "—"
                rows.append(f"• {name} | status: {status} | channel: {channel} | treatment: {treatment}")
            return f"{len(patients)} customer(s):\n" + "\n".join(rows)
    except Exception as e:
        return f"Error querying customers: {e}"


def _tool_get_metrics(settings: Settings) -> str:
    from .db import get_session
    from .models import Message, EscalationEvent
    try:
        with get_session(settings) as session:
            pending = session.query(Message).filter(Message.status == "pending_approval").count()
            sent = session.query(Message).filter(Message.status == "sent").count()
            escalations = session.query(EscalationEvent).filter(
                EscalationEvent.resolved == False  # noqa: E712
            ).count()
            total_customers = session.execute(
                __import__("sqlalchemy").text("SELECT COUNT(*) FROM patients")
            ).scalar()
        return (
            f"Business metrics:\n"
            f"• {pending} messages pending approval in the queue\n"
            f"• {sent} messages sent total\n"
            f"• {escalations} open escalations (need human attention)\n"
            f"• {total_customers} customers in the database"
        )
    except Exception as e:
        return f"Error fetching metrics: {e}"


def _tool_draft_to_queue(inputs: dict, settings: Settings) -> str:
    from .db import get_session
    from .models import Patient, Campaign, Message, CampaignStatus
    customer_name = inputs.get("customer_name", "").strip()
    channel = inputs.get("channel", "sms")
    subject = inputs.get("subject", "")
    body = inputs.get("body", "")
    if not body:
        return "Error: message body is required."
    try:
        with get_session(settings) as session:
            # Find the patient by name (case-insensitive partial match)
            parts = customer_name.split()
            first = parts[0] if parts else ""
            last = parts[-1] if len(parts) > 1 else ""
            patient = (
                session.query(Patient)
                .filter(
                    Patient.first_name.ilike(f"%{first}%"),
                    Patient.last_name.ilike(f"%{last}%") if last else True,
                )
                .first()
            )
            if not patient:
                # Try just first name
                patient = session.query(Patient).filter(
                    Patient.first_name.ilike(f"%{customer_name}%")
                ).first()
            if not patient:
                return (
                    f"Could not find a customer named '{customer_name}' in the database. "
                    f"Use get_customers to see available names first."
                )

            # Find or create a campaign
            campaign = (
                session.query(Campaign)
                .filter(
                    Campaign.patient_id == patient.id,
                    Campaign.status == CampaignStatus.active.value,
                )
                .first()
            )
            if not campaign:
                # Get practice_id from any existing campaign or use default
                existing = session.query(Campaign).first()
                practice_id = existing.practice_id if existing else "default"
                campaign = Campaign(
                    practice_id=practice_id,
                    workflow_id="ai-team-draft",
                    patient_id=patient.id,
                    status=CampaignStatus.active.value,
                )
                session.add(campaign)
                session.flush()

            msg = Message(
                campaign_id=campaign.id,
                direction="outbound",
                channel=channel,
                subject=subject or None,
                body=body,
                status="pending_approval",
                metadata_json={"source": "ai-team"},
            )
            session.add(msg)

        full_name = f"{patient.first_name} {patient.last_name}"
        return (
            f"Message drafted for {full_name} via {channel}. "
            f"It's now in the approval queue — the owner will review it before it sends."
        )
    except Exception as e:
        return f"Error drafting message: {e}"


def _tool_create_document(inputs: dict) -> str:
    title = inputs.get("title", "Document")
    content = inputs.get("content", "")
    if not content:
        return "Error: content is required."
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        import re

        doc = DocxDocument()
        doc.add_heading(title, 0)

        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            # Markdown heading detection
            m = re.match(r"^(#{1,3})\s+(.+)$", stripped)
            if m:
                level = len(m.group(1))
                doc.add_heading(m.group(2), level)
            elif stripped.startswith(("- ", "* ", "• ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            else:
                doc.add_paragraph(stripped)

        out_dir = Path(os.environ.get("ADAPIX_VAR", ".")) / "agent_documents"
        out_dir.mkdir(parents=True, exist_ok=True)
        safe = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")[:50]
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        path = out_dir / f"{safe}_{ts}.docx"
        doc.save(str(path))
        return f"Document saved: {path.name}. The owner can download it from the AI Team chat."
    except Exception as e:
        return f"Error creating document: {e}"


# ---------------------------------------------------------------------------
# Agent loader
# ---------------------------------------------------------------------------
@dataclass
class TeamAgent:
    slug: str
    name: str
    description: str
    emoji: str
    category: str
    system_prompt: str


def _parse_frontmatter(text: str) -> tuple[dict[str, str], str]:
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    fm_text = text[3:end].strip()
    body = text[end + 4:].strip()
    fields: dict[str, str] = {}
    for line in fm_text.splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            fields[k.strip()] = v.strip().strip('"').strip("'")
    return fields, body


def _find_agent_file(slug: str, repo: Path) -> Path | None:
    for md in repo.rglob(f"{slug}.md"):
        return md
    return None


def list_agents() -> list[TeamAgent]:
    repo = _agents_repo()
    agents: list[TeamAgent] = []
    for slug, category in CURATED.items():
        f = _find_agent_file(slug, repo)
        if not f or not f.exists():
            continue
        try:
            text = f.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        fm, body = _parse_frontmatter(text)
        if not fm.get("name"):
            continue
        agents.append(TeamAgent(
            slug=slug,
            name=fm.get("name", slug),
            description=fm.get("description", ""),
            emoji=fm.get("emoji", "◈"),
            category=category,
            system_prompt=body,
        ))
    cat_idx = {c: i for i, c in enumerate(CATEGORY_ORDER)}
    agents.sort(key=lambda a: (cat_idx.get(a.category, 99), a.name))
    return agents


def get_agent(slug: str) -> TeamAgent | None:
    for a in list_agents():
        if a.slug == slug:
            return a
    return None


# ---------------------------------------------------------------------------
# Per-agent chat history
# ---------------------------------------------------------------------------
def _agent_chat_path(slug: str) -> Path:
    var = os.environ.get("ADAPIX_VAR", ".")
    return Path(var) / f"agent_chat_{slug}.json"


def load_agent_history(slug: str) -> list[dict]:
    p = _agent_chat_path(slug)
    if not p.exists():
        return []
    try:
        return json.loads(p.read_text()).get("messages", [])
    except Exception:
        return []


def save_agent_history(slug: str, messages: list[dict]) -> None:
    p = _agent_chat_path(slug)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps({"messages": messages}, indent=2))


def clear_agent_history(slug: str) -> None:
    p = _agent_chat_path(slug)
    if p.exists():
        p.unlink()


# ---------------------------------------------------------------------------
# Chat with tool-use loop
# ---------------------------------------------------------------------------
def _build_system_prompt(agent: TeamAgent) -> str:
    try:
        profile = load_profile()
        context = (
            f"Business name: {profile.practice_name}\n"
            f"Owner: {profile.doctor}\n"
            f"Business type: {profile.practice_type_label or profile.practice_type or 'small business'}\n"
            f"Communication tone: {profile.tone}"
        )
    except Exception:
        context = "Business context not available."

    return (
        f"## TOOL RULES — follow these exactly\n\n"
        f"You have live tool access. The rules below are MANDATORY — they override your default behavior:\n\n"
        f"RULE 1 — search_web: Any time the user says 'search', 'look up', 'find', 'what are the latest', "
        f"'current stats', '2025', or any research request → your FIRST action MUST be to call search_web. "
        f"Never answer a research question from memory alone when search_web is available.\n\n"
        f"RULE 2 — get_customers: Any time the user asks who's in the system, references a specific customer, "
        f"or asks about follow-up/outreach → call get_customers first.\n\n"
        f"RULE 3 — get_metrics: Any question about the business's numbers, performance, pipeline, or queue "
        f"→ call get_metrics.\n\n"
        f"RULE 4 — draft_to_queue: Any request to write, send, or draft a message for a customer "
        f"→ call get_customers to find the name, then call draft_to_queue.\n\n"
        f"RULE 5 — create_document: Any request for a report, strategy document, plan, template, or checklist "
        f"→ call create_document so the owner has a downloadable file.\n\n"
        f"## BUSINESS CONTEXT\n"
        f"{context}\n\n"
        f"## YOUR ROLE\n"
        f"{agent.system_prompt}\n\n"
        f"Be specific to their situation. Keep responses concise and actionable."
    )


def _build_user_content(user_text: str, attachments: list | None) -> str | list:
    """Return a string (no files) or a list of content blocks (with images/text files)."""
    if not attachments:
        return user_text
    blocks: list = []
    if user_text:
        blocks.append({"type": "text", "text": user_text})
    for att in attachments:
        if att["type"] == "image":
            blocks.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": att["media_type"],
                    "data": att["data"],
                },
            })
            blocks.append({"type": "text", "text": f"[Attached image: {att['name']}]"})
        else:
            blocks.append({
                "type": "text",
                "text": f"[Attached file: {att['name']}]\n```\n{att.get('content', '')}\n```",
            })
    return blocks


def send_agent_message(slug: str, user_text: str, attachments: list | None = None) -> dict:
    """Run the tool-use loop and return {reply, tools_used, document_path}."""
    agent = get_agent(slug)
    if not agent:
        raise ValueError(f"Agent not found: {slug}")

    settings = Settings()
    client = Anthropic(api_key=settings.anthropic_api_key)
    history = load_agent_history(slug)
    system = _build_system_prompt(agent)

    messages: list[dict] = [{"role": m["role"], "content": m["content"]} for m in history]
    messages.append({"role": "user", "content": _build_user_content(user_text, attachments)})

    tools_used: list[str] = []
    document_path: str | None = None
    max_iterations = 8

    for _ in range(max_iterations):
        resp = client.messages.create(
            model=settings.adapix_model,
            max_tokens=2048,
            system=system,
            tools=AGENT_TOOLS,
            messages=messages,
        )

        # Collect all text and tool_use blocks
        tool_calls = [b for b in resp.content if b.type == "tool_use"]
        text_blocks = [b for b in resp.content if b.type == "text"]

        if resp.stop_reason == "end_turn" or not tool_calls:
            # Final text response
            reply = "".join(b.text for b in text_blocks if hasattr(b, "text"))
            break

        # Append the assistant's turn (may contain text + tool_use blocks)
        messages.append({"role": "assistant", "content": resp.content})

        # Execute each tool call and collect results
        tool_results = []
        for tc in tool_calls:
            label = TOOL_LABELS.get(tc.name, tc.name)
            if label not in tools_used:
                tools_used.append(label)
            result = _execute_tool(tc.name, tc.input, settings)
            # Track document path if a doc was created
            if tc.name == "create_document" and "Document saved:" in result:
                import re
                m = re.search(r"Document saved: (.+?)\.", result)
                if m:
                    document_path = m.group(1)
            tool_results.append({
                "type": "tool_result",
                "tool_use_id": tc.id,
                "content": result,
            })

        messages.append({"role": "user", "content": tool_results})

    else:
        reply = "I ran into an issue completing that — please try again."

    # Persist to history (store just role + content as text for simplicity)
    now = datetime.utcnow().isoformat() + "Z"
    history.append({"role": "user",      "content": user_text, "ts": now})
    history.append({"role": "assistant", "content": reply,     "ts": now,
                    "tools_used": tools_used})
    save_agent_history(slug, history)

    return {"reply": reply, "tools_used": tools_used, "document_path": document_path}
