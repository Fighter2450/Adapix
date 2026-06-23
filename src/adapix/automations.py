"""Automation engine.

Runs scheduled browser automations: visits a URL, extracts data using Claude,
and saves the result as a Word doc (or plain text / JSON).

Each Automation row in the DB defines:
  - url          — the page to visit
  - task         — natural-language description of what to extract
  - schedule     — cron expression (e.g. "0 9 * * *" = every day at 9am)
  - output_format — docx | txt | json

The engine uses Playwright (headless Chromium) for rendering and Claude for
understanding the page content. No CSS selectors required — Claude reads the
page like a human would and pulls out exactly what the task asks for.
"""
from __future__ import annotations

import json
import logging
import os
from datetime import datetime, UTC
from pathlib import Path
from typing import Any

log = logging.getLogger("adapix.automations")

RESULTS_DIR = Path(__file__).parent.parent.parent / "automation_results"


def _results_dir() -> Path:
    RESULTS_DIR.mkdir(exist_ok=True)
    return RESULTS_DIR


# ---------------------------------------------------------------------------
# Browser + extraction
# ---------------------------------------------------------------------------

def fetch_page_text(
    url: str,
    *,
    login_url: str | None = None,
    login_username: str | None = None,
    login_email: str | None = None,
    login_password: str | None = None,
) -> str:
    """Use Playwright to load the page and return its visible text content.
    If login credentials are provided, navigates to the login page first,
    fills the username/email and password fields, submits, then goes to the
    target URL."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()

        # Log in if credentials provided
        if login_password and (login_username or login_email):
            login_dest = login_url or url
            log.info(f"Logging in at {login_dest}")
            page.goto(login_dest, wait_until="domcontentloaded", timeout=30_000)
            page.wait_for_timeout(1500)

            identifier = login_username or login_email or ""
            # Try common username/email field selectors
            for sel in ['input[type="email"]', 'input[name="email"]',
                        'input[name="username"]', 'input[name="user"]',
                        'input[type="text"]']:
                if page.locator(sel).count() > 0:
                    page.locator(sel).first.fill(identifier)
                    break

            # Fill password
            for sel in ['input[type="password"]', 'input[name="password"]']:
                if page.locator(sel).count() > 0:
                    page.locator(sel).first.fill(login_password)
                    break

            # Submit — try button first, then Enter
            submitted = False
            for sel in ['button[type="submit"]', 'input[type="submit"]',
                        'button:has-text("Log in")', 'button:has-text("Sign in")',
                        'button:has-text("Login")', 'button:has-text("Continue")']:
                if page.locator(sel).count() > 0:
                    page.locator(sel).first.click()
                    submitted = True
                    break
            if not submitted:
                page.keyboard.press("Enter")

            page.wait_for_load_state("domcontentloaded", timeout=15_000)
            page.wait_for_timeout(2000)

            # Navigate to target if different from login page
            if url != login_dest:
                page.goto(url, wait_until="domcontentloaded", timeout=30_000)
                page.wait_for_timeout(2000)
        else:
            page.goto(url, wait_until="domcontentloaded", timeout=30_000)
            page.wait_for_timeout(2000)

        text = page.evaluate("""() => {
            ['script','style','nav','footer','header'].forEach(tag => {
                document.querySelectorAll(tag).forEach(el => el.remove());
            });
            return document.body.innerText;
        }""")
        browser.close()
    return text or ""


def _get_page_snapshot(page) -> tuple[str, list[dict]]:
    """Return (visible_text, links) from a Playwright page."""
    text = page.evaluate("""() => {
        ['script','style','nav','footer','header'].forEach(t =>
            document.querySelectorAll(t).forEach(el => el.remove()));
        return document.body.innerText;
    }""") or ""
    links = page.evaluate("""() => {
        const seen = new Set();
        return Array.from(document.querySelectorAll('a[href]'))
            .map(a => ({ text: a.innerText.trim().slice(0, 80), href: a.href }))
            .filter(l => l.href.startsWith('http') && l.text && !seen.has(l.href) && seen.add(l.href))
            .slice(0, 60);
    }""") or []
    return text, links


def _decide_next_step(client, model: str, task: str, current_url: str,
                      page_text: str, links: list[dict], step: int) -> dict:
    """Ask Claude whether to extract from this page or navigate somewhere else.
    Returns {'action': 'extract', 'data': '...'} or {'action': 'navigate', 'url': '...'}."""
    links_block = "\n".join(f"- {l['text']}: {l['href']}" for l in links[:40]) or "(no links found)"
    prompt = f"""You are an AI browser agent. A user wants you to: {task}

You are currently on: {current_url}
This is step {step + 1} of up to 5.

PAGE CONTENT (first 6000 chars):
{page_text[:6000]}

LINKS ON THIS PAGE:
{links_block}

Decide what to do:
1. If this page already contains the data needed to complete the task, respond with:
   ACTION: extract
   DATA: [the extracted data, well-formatted]

2. If you need to navigate to a different page to find the data, respond with:
   ACTION: navigate
   URL: [the full URL to navigate to — must be one of the links listed above or a logical variation of the current URL]
   REASON: [one sentence why]

Respond in exactly that format. Do not add any other text before ACTION:."""

    msg = client.messages.create(
        model=model,
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}],
    )
    text = msg.content[0].text.strip()

    if text.startswith("ACTION: extract") or "ACTION: extract" in text:
        data_marker = "DATA:"
        idx = text.find(data_marker)
        data = text[idx + len(data_marker):].strip() if idx >= 0 else text
        return {"action": "extract", "data": data}

    if "ACTION: navigate" in text:
        url_marker = "URL:"
        idx = text.find(url_marker)
        url_line = text[idx + len(url_marker):].split("\n")[0].strip() if idx >= 0 else ""
        if url_line:
            return {"action": "navigate", "url": url_line}

    # Fallback: if Claude didn't follow the format, just extract what's there
    return {"action": "extract", "data": text}


def navigate_and_extract(
    url: str,
    task: str,
    *,
    login_url: str | None = None,
    login_username: str | None = None,
    login_email: str | None = None,
    login_password: str | None = None,
    max_steps: int = 5,
) -> tuple[str, str]:
    """Open a browser, navigate up to max_steps pages to find the data,
    and return (final_url, extracted_data)."""
    import anthropic
    from playwright.sync_api import sync_playwright
    from .config import Settings

    settings = Settings()
    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()

        # Handle login
        if login_password and (login_username or login_email):
            login_dest = login_url or url
            log.info(f"Logging in at {login_dest}")
            page.goto(login_dest, wait_until="domcontentloaded", timeout=30_000)
            page.wait_for_timeout(1500)
            identifier = login_username or login_email or ""
            for sel in ['input[type="email"]', 'input[name="email"]',
                        'input[name="username"]', 'input[type="text"]']:
                if page.locator(sel).count() > 0:
                    page.locator(sel).first.fill(identifier)
                    break
            for sel in ['input[type="password"]', 'input[name="password"]']:
                if page.locator(sel).count() > 0:
                    page.locator(sel).first.fill(login_password)
                    break
            submitted = False
            for sel in ['button[type="submit"]', 'input[type="submit"]',
                        'button:has-text("Log in")', 'button:has-text("Sign in")']:
                if page.locator(sel).count() > 0:
                    page.locator(sel).first.click()
                    submitted = True
                    break
            if not submitted:
                page.keyboard.press("Enter")
            page.wait_for_load_state("domcontentloaded", timeout=15_000)
            page.wait_for_timeout(2000)
            if url != login_dest:
                page.goto(url, wait_until="domcontentloaded", timeout=30_000)
                page.wait_for_timeout(2000)
        else:
            page.goto(url, wait_until="domcontentloaded", timeout=30_000)
            page.wait_for_timeout(2000)

        final_url = page.url
        extracted = ""

        for step in range(max_steps):
            log.info(f"Agent step {step + 1}: {page.url}")
            page_text, links = _get_page_snapshot(page)
            decision = _decide_next_step(client, settings.adapix_model,
                                         task, page.url, page_text, links, step)

            if decision["action"] == "extract":
                final_url = page.url
                extracted = decision["data"]
                log.info(f"Agent extracted data at step {step + 1} from {final_url}")
                break

            nav_url = decision.get("url", "")
            if not nav_url:
                # No URL to navigate to — extract from current page
                extracted = page_text[:8000]
                break
            log.info(f"Agent navigating to: {nav_url}")
            try:
                page.goto(nav_url, wait_until="domcontentloaded", timeout=30_000)
                page.wait_for_timeout(2000)
            except Exception as nav_err:
                log.warning(f"Navigation failed: {nav_err} — extracting from current page")
                page_text, _ = _get_page_snapshot(page)
                extracted = page_text[:8000]
                break
        else:
            # Hit max_steps — extract whatever is on screen
            final_url = page.url
            page_text, _ = _get_page_snapshot(page)
            extracted = page_text[:8000]

        browser.close()

    return final_url, extracted


# ---------------------------------------------------------------------------
# Output file generation
# ---------------------------------------------------------------------------

def save_as_docx(automation_name: str, url: str, task: str, data: str, run_id: int) -> str:
    """Save extracted data as a Word document. Returns the file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(automation_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n").bold = False
    meta.add_run(f"Source: {url}\n").bold = False
    meta.add_run(f"Task: {task}").bold = False
    meta.style = doc.styles["Normal"]

    doc.add_paragraph()  # spacer

    # Data
    doc.add_heading("Extracted Data", level=2)
    for line in data.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)

    out_dir = _results_dir()
    filename = f"automation_{run_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = str(out_dir / filename)
    doc.save(path)
    return path


def save_as_txt(automation_name: str, url: str, task: str, data: str, run_id: int) -> str:
    out_dir = _results_dir()
    filename = f"automation_{run_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    path = str(out_dir / filename)
    content = f"{automation_name}\n{'='*len(automation_name)}\nGenerated: {datetime.now()}\nSource: {url}\nTask: {task}\n\n{data}"
    Path(path).write_text(content, encoding="utf-8")
    return path


def save_as_json(automation_name: str, url: str, task: str, data: str, run_id: int) -> str:
    out_dir = _results_dir()
    filename = f"automation_{run_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    path = str(out_dir / filename)
    payload = {
        "automation": automation_name,
        "url": url,
        "task": task,
        "generated_at": datetime.now().isoformat(),
        "data": data,
    }
    Path(path).write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# Run one automation
# ---------------------------------------------------------------------------

def run_automation(automation_id: int) -> dict[str, Any]:
    """Execute a single automation. Creates an AutomationRun row, runs the browser,
    calls Claude, saves the output file, and updates the Automation + Run rows."""
    from .config import Settings
    from .db import get_session
    from .models import Automation, AutomationRun

    settings = Settings()

    with get_session(settings) as s:
        auto = s.get(Automation, automation_id)
        if not auto:
            return {"ok": False, "error": f"Automation {automation_id} not found"}

        run = AutomationRun(automation_id=automation_id, status="running")
        s.add(run)
        s.flush()
        run_id = run.id

        name = auto.name
        url = auto.url
        task = auto.task
        fmt = auto.output_format or "docx"
        login_url = auto.login_url
        login_username = auto.login_username
        login_email = auto.login_email
        login_password = auto.login_password

    try:
        log.info(f"Automation {automation_id} ({name}): starting agent at {url}")
        final_url, data = navigate_and_extract(
            url,
            task,
            login_url=login_url,
            login_username=login_username,
            login_email=login_email,
            login_password=login_password,
        )

        log.info(f"Automation {automation_id}: saving {fmt}")
        if fmt == "txt":
            path = save_as_txt(name, final_url, task, data, run_id)
        elif fmt == "json":
            path = save_as_json(name, final_url, task, data, run_id)
        else:
            path = save_as_docx(name, final_url, task, data, run_id)

        now = datetime.now(UTC).replace(tzinfo=None)
        with get_session(settings) as s:
            run = s.get(AutomationRun, run_id)
            run.status = "ok"
            run.finished_at = now
            run.extracted_data = data
            run.result_path = path
            auto = s.get(Automation, automation_id)
            auto.last_run_at = now
            auto.last_run_status = "ok"
            auto.last_result_path = path
            auto.last_error = None

        log.info(f"Automation {automation_id}: done → {path}")
        return {"ok": True, "run_id": run_id, "path": path, "data_preview": data[:300]}

    except Exception as exc:
        now = datetime.now(UTC).replace(tzinfo=None)
        err = str(exc)
        log.error(f"Automation {automation_id} failed: {err}")
        with get_session(settings) as s:
            run = s.get(AutomationRun, run_id)
            run.status = "error"
            run.finished_at = now
            run.error = err
            auto = s.get(Automation, automation_id)
            auto.last_run_at = now
            auto.last_run_status = "error"
            auto.last_error = err
        return {"ok": False, "error": err}


# ---------------------------------------------------------------------------
# Scheduler — called by main.py background loop
# ---------------------------------------------------------------------------

def get_due_automations() -> list[int]:
    """Return IDs of active automations whose cron schedule is currently due."""
    from croniter import croniter
    from .config import Settings
    from .db import get_session
    from .models import Automation

    settings = Settings()
    due = []
    now = datetime.now(UTC).replace(tzinfo=None)

    with get_session(settings) as s:
        autos = s.query(Automation).filter(Automation.status == "active").all()
        for a in autos:
            try:
                cron = croniter(a.schedule)
                prev = cron.get_prev(datetime)
                # Due if last scheduled time is within the past 6 minutes
                # (scheduler runs every 5 min — 6 min window avoids drift misses)
                minutes_ago = (now - prev).total_seconds() / 60
                last_ran = a.last_run_at
                already_ran = last_ran and (now - last_ran).total_seconds() < 360
                if minutes_ago <= 6 and not already_ran:
                    due.append(a.id)
            except Exception as exc:
                log.warning(f"Automation {a.id} bad cron '{a.schedule}': {exc}")

    return due
